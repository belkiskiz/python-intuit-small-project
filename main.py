# This is a sample Python script.

# Press ⌃R to execute it or replace it with your code.
# Press Double ⇧ to search everywhere for classes, files, tool windows, actions, and settings.


from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)



document = Document()

# profile picture
document.add_picture(
    'images.jpeg', width=Inches(2.0)
)

#name phone number and email details
name = input('what is your name ? ')
speak('Hello' + name + 'How are you today ?')
phone_number = input('what is your phone number ? ')
speak('what is your phone number ? ')
email = 'just_stuff@stuff.com'

document.add_paragraph(name+'%'+phone_number+'/'+email)
# about me
document.add_heading('About me')
about_me = input('Tell me about yourself')
document.add_paragraph(about_me)

# work experience
document.add_heading('Work Experience')

p = document.add_paragraph()

company = input('Enter company')
start_date = input('From date')
end_date = input('To date')

p.add_run(company + ' ').bold = True
p.add_run(start_date + '-' + end_date + '\n').italic = True
experience_details = input('Describe your experience at' + company)
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input('Do you have more experiences ? Yes or No')
    if has_more_experiences.lower() =='yes':

        p = document.add_paragraph()

        company = input('Enter company')
        start_date = input('From date')
        end_date = input('To date')

        p.add_run(company + ' ').bold = True
        p.add_run(start_date + '-' + end_date + '\n').italic = True
        experience_details = input('Describe your experience at' + company)
        p.add_run(experience_details)

    else:
        break

# skills
document.add_heading('Skills ')
skill = input('What is your skill ? ')
g = document.add_paragraph(skill)
g.style = 'List Bullet'
while True:
    has_more_skills = input('Do you have any other skills ? ')
    if has_more_skills.lower() == 'yes':
        skill = input('What is your skill ? ')
        g = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using Amigoscode'





document.save('cv.docx')

